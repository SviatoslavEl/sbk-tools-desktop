from __future__ import annotations

import hashlib
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pypdfium2 as pdfium
from docx import Document


MAX_SOURCE_BYTES = 250 * 1024 * 1024
MAX_EXTRACTED_CHARS = 20 * 1024 * 1024
MAX_ZIP_UNPACKED_BYTES = 500 * 1024 * 1024
MAX_FRAGMENTS = 50_000
NS_MAIN = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
NS_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_PACKAGE_REL = "http://schemas.openxmlformats.org/package/2006/relationships"


def _sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def _validate_source(path: Path) -> tuple[int, bytes]:
    if not path.is_file():
        raise ValueError("Документ не найден.")
    size = path.stat().st_size
    if size <= 0 or size > MAX_SOURCE_BYTES:
        raise ValueError("Размер документа должен быть от 1 байта до 250 МБ.")
    with path.open("rb") as stream:
        magic = stream.read(8)
    return size, magic


def _bounded(fragments: list[dict]) -> tuple[list[dict], str]:
    if len(fragments) > MAX_FRAGMENTS:
        raise ValueError("Документ содержит слишком много текстовых фрагментов.")
    total = 0
    texts: list[str] = []
    for fragment in fragments:
        text = str(fragment.get("text", "")).strip()
        total += len(text)
        if total > MAX_EXTRACTED_CHARS:
            raise ValueError("Извлечённый текст превышает безопасный предел 20 МБ.")
        if text:
            texts.append(text)
    return fragments, "\n".join(texts)


def _pdf(path: Path) -> tuple[str, list[dict], list[str]]:
    document = pdfium.PdfDocument(path)
    fragments: list[dict] = []
    try:
        for index in range(len(document)):
            page = document[index]
            text_page = page.get_textpage()
            try:
                text = text_page.get_text_range().strip()
            finally:
                text_page.close()
                page.close()
            fragments.append({"id": f"page-{index + 1}", "locator": f"Страница {index + 1}", "page": index + 1, "text": text})
    finally:
        document.close()
    warnings = [] if any(entry["text"] for entry in fragments) else ["В PDF нет текстового слоя. Для поиска сначала выполните OCR в модуле сканирования."]
    return "application/pdf", fragments, warnings


def _docx(path: Path) -> tuple[str, list[dict], list[str]]:
    with zipfile.ZipFile(path) as archive:
        total = sum(entry.file_size for entry in archive.infolist())
        if total > MAX_ZIP_UNPACKED_BYTES:
            raise ValueError("Распакованный DOCX превышает безопасный предел.")
        names = set(archive.namelist())
        if "word/document.xml" not in names or "[Content_Types].xml" not in names:
            raise ValueError("DOCX повреждён или имеет неподдерживаемую структуру.")
        if any(name.startswith("../") or "/../" in name for name in names):
            raise ValueError("DOCX содержит опасные пути.")
    document = Document(str(path))
    fragments: list[dict] = []
    for index, paragraph in enumerate(document.paragraphs, 1):
        if paragraph.text.strip():
            fragments.append({"id": f"paragraph-{index}", "locator": f"Абзац {index}", "section": f"Абзац {index}", "text": paragraph.text})
    for table_index, table in enumerate(document.tables, 1):
        for row_index, row in enumerate(table.rows, 1):
            text = "\t".join(cell.text.strip() for cell in row.cells)
            if text.strip():
                fragments.append({"id": f"table-{table_index}-row-{row_index}", "locator": f"Таблица {table_index}, строка {row_index}", "section": f"Таблица {table_index}", "text": text})
    for section_index, section in enumerate(document.sections, 1):
        for kind, paragraphs in (("Верхний колонтитул", section.header.paragraphs), ("Нижний колонтитул", section.footer.paragraphs)):
            text = "\n".join(paragraph.text.strip() for paragraph in paragraphs if paragraph.text.strip())
            if text:
                fragments.append({"id": f"section-{section_index}-{kind}", "locator": f"Раздел {section_index}: {kind}", "section": f"Раздел {section_index}", "text": text})
    return "application/vnd.openxmlformats-officedocument.wordprocessingml.document", fragments, []


def _xlsx(path: Path) -> tuple[str, list[dict], list[str]]:
    fragments: list[dict] = []
    warnings: list[str] = []
    with zipfile.ZipFile(path) as archive:
        total = sum(entry.file_size for entry in archive.infolist())
        if total > MAX_ZIP_UNPACKED_BYTES:
            raise ValueError("Распакованный XLSX превышает безопасный предел.")
        names = set(archive.namelist())
        if "xl/workbook.xml" not in names or "[Content_Types].xml" not in names:
            raise ValueError("XLSX повреждён или имеет неподдерживаемую структуру.")
        shared: list[str] = []
        if "xl/sharedStrings.xml" in names:
            root = ET.fromstring(archive.read("xl/sharedStrings.xml"))
            shared = ["".join(node.text or "" for node in item.iter(f"{{{NS_MAIN}}}t")) for item in root.findall(f"{{{NS_MAIN}}}si")]
        relations = ET.fromstring(archive.read("xl/_rels/workbook.xml.rels"))
        targets = {entry.attrib["Id"]: entry.attrib["Target"] for entry in relations.findall(f"{{{NS_PACKAGE_REL}}}Relationship")}
        workbook = ET.fromstring(archive.read("xl/workbook.xml"))
        for sheet in workbook.findall(f".//{{{NS_MAIN}}}sheet"):
            name = sheet.attrib.get("name", "Лист")
            if sheet.attrib.get("state", "visible") != "visible":
                warnings.append(f"Лист «{name}» скрыт и включён в проверку.")
            relation_id = sheet.attrib.get(f"{{{NS_REL}}}id", "")
            target = targets.get(relation_id, "")
            sheet_path = target.lstrip("/") if target.startswith("/xl/") else f"xl/{target.lstrip('/')}"
            if sheet_path not in names:
                continue
            root = ET.fromstring(archive.read(sheet_path))
            for row in root.findall(f".//{{{NS_MAIN}}}row"):
                values: list[str] = []
                first_cell = ""
                last_cell = ""
                for cell in row.findall(f"{{{NS_MAIN}}}c"):
                    reference = cell.attrib.get("r", "")
                    first_cell = first_cell or reference
                    last_cell = reference
                    kind = cell.attrib.get("t")
                    if kind == "inlineStr":
                        value = "".join(node.text or "" for node in cell.iter(f"{{{NS_MAIN}}}t"))
                    else:
                        node = cell.find(f"{{{NS_MAIN}}}v")
                        value = node.text if node is not None and node.text is not None else ""
                        if kind == "s" and value.isdigit() and int(value) < len(shared):
                            value = shared[int(value)]
                    values.append(value)
                text = "\t".join(values).strip()
                if text:
                    cell_range = first_cell if first_cell == last_cell else f"{first_cell}:{last_cell}"
                    fragments.append({"id": f"sheet-{len(fragments) + 1}", "locator": f"Лист {name}, {cell_range}", "sheet": name, "cellRange": cell_range, "text": text})
    return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", fragments, warnings


def extract_document(path: Path) -> dict:
    size, magic = _validate_source(path)
    suffix = path.suffix.lower()
    if magic.startswith(b"%PDF-"):
        mime_type, fragments, warnings = _pdf(path)
    elif magic.startswith(b"PK\x03\x04") and suffix == ".docx":
        mime_type, fragments, warnings = _docx(path)
    elif magic.startswith(b"PK\x03\x04") and suffix == ".xlsx":
        mime_type, fragments, warnings = _xlsx(path)
    else:
        raise ValueError("Поддерживаются только настоящие PDF, DOCX и XLSX.")
    fragments, text = _bounded(fragments)
    return {"type": "extraction", "fileName": path.name, "mimeType": mime_type, "sizeBytes": size, "sha256": _sha256(path), "extractedText": text, "fragments": fragments, "warnings": warnings, "extractionEngineVersion": "sbk-local-extractor/1", "protocolVersion": 2}
