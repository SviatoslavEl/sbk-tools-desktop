from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "demo-data"
FONT_REGULAR = ROOT / "scanner-worker/resources/fonts/NotoSans-Regular.ttf"
FONT_BOLD = ROOT / "scanner-worker/resources/fonts/NotoSans-Bold.ttf"
GREEN = "2F9845"
DARK = "173C2A"
LIGHT = "DFF1E3"


def shade(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    fill = OxmlElement("w:shd")
    fill.set(qn("w:fill"), color)
    properties.append(fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        margins.append(node)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Cm(widths[index])
        cell.text = header
        shade(cell, GREEN)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9)
        set_cell_margins(cell)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].width = Cm(widths[index])
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[index])
            if row_index % 2:
                shade(cells[index], "F2F8F3")
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)


def build_docx() -> Path:
    output = OUTPUT / "Демо_закупка_лабораторного_оборудования.docx"
    document = Document()
    section = document.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    section.left_margin = section.right_margin = Cm(2)
    section.top_margin, section.bottom_margin = Cm(1.8), Cm(1.8)
    section.header_distance, section.footer_distance = Cm(0.8), Cm(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Noto Sans"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for name, size in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13)):
        style = styles[name]
        style.font.name = "Noto Sans"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(GREEN if name != "Title" else DARK)
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(7)

    header = section.header.paragraphs[0]
    header.text = "СБК ИНСТРУМЕНТЫ  /  ДЕМО-ДОКУМЕНТ"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Noto Sans"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(GREEN)
        run.bold = True
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Учебные данные - не являются реальной закупкой")
    run.font.name = "Noto Sans"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 119, 107)

    title = document.add_paragraph(style="Title")
    title.add_run("Поставка лабораторного оборудования")
    subtitle = document.add_paragraph("Демонстрационная закупочная документация")
    subtitle.style = styles["Subtitle"]
    subtitle.runs[0].font.color.rgb = RGBColor(80, 114, 90)
    document.add_paragraph(
        "Извещение № ДЕМО-2026-014 · Заказчик: АНО «Учебный центр СБК» · "
        "Начальная цена: 4 850 000 руб. · Срок подачи: 18 сентября 2026 года, 12:00"
    )

    document.add_heading("1. Предмет и состав поставки", level=1)
    document.add_paragraph(
        "Поставщик выполняет поставку, монтаж, первичную настройку и обучение двух специалистов заказчика. "
        "Все наименования и реквизиты созданы специально для демонстрации функций приложения."
    )
    add_table(
        document,
        ["№", "Наименование", "Кол-во", "Цена, руб.", "Срок"],
        [
            ["1", "Спектрофотометр учебный SP-200", "2 шт.", "1 240 000", "45 дней"],
            ["2", "Центрифуга лабораторная CL-24", "4 шт.", "1 520 000", "35 дней"],
            ["3", "Комплект дозаторов DP-8", "10 компл.", "890 000", "25 дней"],
            ["4", "Монтаж, настройка и обучение", "1 усл.", "1 200 000", "50 дней"],
        ],
        [0.8, 7.7, 2.0, 2.8, 2.2],
    )
    document.add_heading("2. Ключевые условия", level=1)
    for text in (
        "Гарантия - не менее 24 месяцев с даты подписания акта.",
        "Оплата - в течение 15 рабочих дней после приёмки.",
        "Обеспечение заявки - 1% от начальной цены.",
        "Требуется подтверждённый опыт двух сопоставимых поставок.",
    ):
        document.add_paragraph(text, style="List Bullet")

    document.add_page_break()
    document.add_heading("3. Матрица соответствия", level=1)
    document.add_paragraph(
        "Страница начата явным разрывом. Она помогает проверить сохранение границ страниц при обработке DOCX."
    )
    add_table(
        document,
        ["Требование", "Предложение участника", "Статус"],
        [
            ["Диапазон спектра 190-1100 нм", "Модель SP-200: 185-1100 нм", "Соответствует"],
            ["Ротор не менее 24 мест", "Ротор 24 x 15 мл включён", "Соответствует"],
            ["Обучение на площадке заказчика", "8 академических часов", "Соответствует"],
            ["Сервисный центр в регионе", "Выезд инженера в течение 2 дней", "Требует подтверждения"],
        ],
        [5.4, 6.7, 3.4],
    )
    document.add_heading("4. Контрольный график", level=1)
    add_table(
        document,
        ["Этап", "Ответственный", "Дата", "Готовность"],
        [
            ["Анализ документации", "Анна Волкова", "28.08.2026", "100%"],
            ["Запрос цен", "Илья Морозов", "02.09.2026", "70%"],
            ["Проверка обеспечения", "Мария Соколова", "08.09.2026", "40%"],
            ["Финальная проверка заявки", "Руководитель группы", "17.09.2026", "0%"],
        ],
        [5.8, 4.5, 2.8, 2.4],
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("5. Декларация участника", level=1)
    document.add_paragraph(
        "ООО «Демо Лаб» подтверждает отсутствие задолженности, конфликта интересов и ограничений, "
        "препятствующих участию в закупке. Сведения в этом разделе являются вымышленными."
    )
    for label, value in (
        ("ИНН", "7700000000"),
        ("КПП", "770001001"),
        ("Адрес", "г. Москва, Учебная улица, д. 10"),
        ("Контакт", "Елена Демова, +7 900 000-00-00"),
    ):
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)
    document.add_paragraph("Подпись уполномоченного лица: ____________________")
    document.add_paragraph("Дата: «___» __________ 2026 г.")
    document.core_properties.title = "Демонстрационная закупка СБК"
    document.core_properties.subject = "Тест DOCX, разрывов страниц, таблиц и кириллицы"
    document.core_properties.author = "СБК Инструменты"
    document.save(output)
    return output


def scan_page(lines: list[tuple[str, int]], page_number: int) -> Image.Image:
    image = Image.new("RGB", (1654, 2339), "#f7f4ec")
    draw = ImageDraw.Draw(image)
    regular = ImageFont.truetype(str(FONT_REGULAR), 34)
    bold = ImageFont.truetype(str(FONT_BOLD), 47)
    small = ImageFont.truetype(str(FONT_REGULAR), 27)
    draw.rectangle((95, 85, 1559, 2250), outline="#b8b1a4", width=3)
    draw.text((132, 125), "КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ", font=bold, fill="#242824")
    draw.text((132, 205), f"Лист {page_number} из 2 · Демо-скан для автономного OCR", font=small, fill="#5c625d")
    y = 315
    for text, indent in lines:
        draw.text((132 + indent, y), text, font=regular, fill="#292d29")
        y += 68
    draw.line((132, y + 25, 1505, y + 25), fill="#777b75", width=2)
    draw.text((132, y + 65), "ООО «Демо Лаб»   ИНН 7700000000", font=small, fill="#3d433d")
    draw.text((1110, 2160), "Подпись __________", font=small, fill="#3d433d")
    return image.rotate(0.18 if page_number == 1 else -0.12, fillcolor="#dedbd3", resample=Image.Resampling.BICUBIC)


def build_scan_files() -> tuple[Path, Path]:
    pages = [
        scan_page([
            ("Заказчику: АНО «Учебный центр СБК»", 0),
            ("Предлагаем поставить оборудование:", 0),
            ("1. Спектрофотометр SP-200 - 2 шт.", 28),
            ("2. Центрифуга CL-24 - 4 шт.", 28),
            ("3. Комплект дозаторов DP-8 - 10 компл.", 28),
            ("Общая стоимость: 4 620 000 рублей.", 0),
            ("Срок поставки: не более 45 календарных дней.", 0),
        ], 1),
        scan_page([
            ("Условия предложения:", 0),
            ("Гарантия на оборудование - 24 месяца.", 28),
            ("Монтаж и обучение включены в стоимость.", 28),
            ("Оплата в течение 15 рабочих дней.", 28),
            ("Предложение действительно до 30.09.2026.", 28),
            ("Контакт: Елена Демова, +7 900 000-00-00", 0),
        ], 2),
    ]
    pdf_path = OUTPUT / "Демо_скан_коммерческого_предложения.pdf"
    pdf = canvas.Canvas(str(pdf_path), pagesize=A4, pageCompression=1)
    for page in pages:
        buffer = io.BytesIO()
        page.save(buffer, format="JPEG", quality=76, optimize=True)
        buffer.seek(0)
        from reportlab.lib.utils import ImageReader

        pdf.drawImage(ImageReader(buffer), 0, 0, width=A4[0], height=A4[1])
        pdf.showPage()
    pdf.save()
    image_path = OUTPUT / "Демо_скан_ценового_предложения.jpg"
    pages[0].save(image_path, format="JPEG", quality=80, optimize=True)
    return pdf_path, image_path


def main() -> None:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_scan_files()


if __name__ == "__main__":
    main()
